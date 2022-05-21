import qrcode
from docx import Document
from docx.shared import Inches
from os import listdir
import datetime as datetime
import random

def listproduce():        #生成一个包含多个随机四位数（不重复）的数组
    list = []

    finlist = []
    for i in range(0, 150):     #生成一个包含多个随机四位数的数组
        randnum = random.randint(1000, 9999)
        list.append(randnum)
    list1 = set(list)
    print(list1)
    for j in list1:    #保证四位数的唯一性
        if j not in finlist:
            finlist.append(j)
    return finlist







# Link for website
images = []
inventoryDate = datetime.datetime.now().strftime("%Y%m%d")    #引入当天日期
rannum = listproduce()
for i in range(0,10):     #生成10个二维码图片
    input_data =  inventoryDate  +  str(rannum[i])

    qr = qrcode.QRCode(
        version=1,
        box_size=10,
        border=5)

    qr.add_data(input_data)
    qr.make(fit=True)

    img = qr.make_image(fill='black', back_color='white')

    img.save(r'D:\\二维码生成\\%s.png'%input_data)

doc = Document()  # doc对象
for fn in listdir(r'D:\\二维码生成'):     #二维码插入word以便打印粘贴再零件上
    print(fn)
    if fn.endswith('.png'):
        b = fn.replace('.png', '', 1)
        doc.add_paragraph(str(b))  # 添加文字
        doc.add_picture(fn, width=Inches(1.5))  # 添加图, 设置宽度
doc.save(r'D:\\二维码生成\\word文档.docx')  # 保存路径
